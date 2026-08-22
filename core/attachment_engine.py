"""
UMAY Attachment Engine

Handles file uploads for chat context:
- Text/code files: direct content extraction
- PDF: text extraction via pypdf
- DOCX: text extraction via python-docx
- XLSX/CSV: data extraction via openpyxl/csv
- Images: base64 encoding for vision models (llava, gemma3)
"""

from __future__ import annotations

import base64
import csv
import hashlib
import io
import json
import os
import tempfile
from datetime import datetime
from pathlib import Path
from typing import Any

# File size limits
MAX_FILE_SIZE = 20 * 1024 * 1024  # 20 MB
MAX_TEXT_CONTENT = 15000  # max chars for text content sent to model
MAX_TABLE_ROWS = 100  # max rows for spreadsheet data

# Supported extensions
TEXT_EXTENSIONS = {
    ".txt", ".md", ".py", ".js", ".ts", ".jsx", ".tsx", ".json",
    ".csv", ".html", ".css", ".xml", ".yaml", ".yml", ".toml",
    ".log", ".sh", ".bat", ".ps1", ".sql", ".r", ".go", ".rs",
    ".java", ".c", ".cpp", ".h", ".hpp", ".rb", ".php", ".swift",
    ".kt", ".scala", ".ini", ".cfg", ".conf", ".env", ".gitignore",
    ".dockerfile", ".makefile",
}

IMAGE_EXTENSIONS = {".jpg", ".jpeg", ".png", ".webp", ".gif", ".bmp", ".tiff"}

PDF_EXTENSIONS = {".pdf"}

DOCX_EXTENSIONS = {".docx"}

XLSX_EXTENSIONS = {".xlsx", ".xls"}

UNSUPPORTED_EXTENSIONS = {".exe", ".dll", ".so", ".dylib", ".bin", ".zip", ".rar", ".7z", ".tar", ".gz"}


def classify_file(filename: str) -> str:
    """Classify file type for processing pipeline."""
    ext = Path(filename).suffix.lower()
    if ext in TEXT_EXTENSIONS:
        return "text"
    if ext in IMAGE_EXTENSIONS:
        return "image"
    if ext in PDF_EXTENSIONS:
        return "pdf"
    if ext in DOCX_EXTENSIONS:
        return "docx"
    if ext in XLSX_EXTENSIONS:
        return "xlsx"
    if ext == ".csv":
        return "csv"
    if ext in UNSUPPORTED_EXTENSIONS:
        return "unsupported"
    return "unknown"


def get_file_icon(filename: str) -> str:
    """Get icon for file type."""
    ext = Path(filename).suffix.lower()
    if ext in IMAGE_EXTENSIONS:
        return "🖼️"
    if ext in PDF_EXTENSIONS:
        return "📄"
    if ext in DOCX_EXTENSIONS:
        return "📝"
    if ext in XLSX_EXTENSIONS or ext == ".csv":
        return "📊"
    if ext in TEXT_EXTENSIONS:
        return "📎"
    return "📁"


def extract_text_content(file_data: bytes, filename: str) -> dict[str, Any]:
    """Extract text content from file based on type."""
    ext = Path(filename).suffix.lower()
    file_type = classify_file(filename)

    result = {
        "filename": filename,
        "ext": ext,
        "type": file_type,
        "icon": get_file_icon(filename),
        "size": len(file_data),
        "content": "",
        "metadata": {},
        "error": None,
    }

    try:
        if file_type == "text":
            result["content"] = _extract_text(file_data, ext)
        elif file_type == "csv":
            result["content"] = _extract_csv(file_data)
        elif file_type == "pdf":
            result["content"] = _extract_pdf(file_data)
            result["metadata"]["pages"] = _count_pdf_pages(file_data)
        elif file_type == "docx":
            result["content"] = _extract_docx(file_data)
        elif file_type == "xlsx":
            result["content"] = _extract_xlsx(file_data, filename)
        elif file_type == "image":
            b64 = base64.b64encode(file_data).decode()
            result["content"] = ""  # No text content
            result["base64"] = b64
            result["metadata"]["is_vision"] = True
        elif file_type == "unsupported":
            result["error"] = f"Bu dosya formatı desteklenmiyor: {ext}"
        else:
            result["error"] = f"Bilinmeyen dosya formatı: {ext}"
    except Exception as e:
        result["error"] = f"Dosya işlenirken hata: {str(e)[:200]}"

    return result


def _extract_text(file_data: bytes, ext: str) -> str:
    """Extract text from plain text / code files."""
    try:
        text = file_data.decode("utf-8", errors="replace")
    except Exception:
        text = file_data.decode("latin-1", errors="replace")

    lines = text.splitlines()
    if len(lines) > 300:
        # For very large files, include beginning + end + summary
        header = "\n".join(lines[:100])
        footer = "\n".join(lines[-50:])
        content = f"--- Dosya: {len(lines)} satır (ilk 100 + son 50 gösteriliyor) ---\n\n{header}\n\n...\n\n{footer}"
    else:
        content = text

    return content[:MAX_TEXT_CONTENT]


def _extract_csv(file_data: bytes) -> str:
    """Extract CSV content as readable text."""
    try:
        text = file_data.decode("utf-8", errors="replace")
        reader = csv.reader(io.StringIO(text))
        rows = list(reader)

        if not rows:
            return "Boş CSV dosyası."

        # Header + first rows
        header = rows[0]
        data_rows = rows[1:MAX_TABLE_ROWS + 1]
        total_rows = len(rows) - 1

        lines = [f"Sütunlar: {', '.join(header)}"]
        lines.append(f"Toplam satır: {total_rows}")
        lines.append("")

        for i, row in enumerate(data_rows[:30]):  # Show first 30 rows
            lines.append(f"  Satır {i+1}: {', '.join(str(c) for c in row)}")

        if total_rows > 30:
            lines.append(f"\n... ve {total_rows - 30} satır daha.")

        return "\n".join(lines)
    except Exception as e:
        return f"CSV okuma hatası: {e}"


def _extract_pdf(file_data: bytes) -> str:
    """Extract text from PDF using pypdf."""
    try:
        from pypdf import PdfReader

        reader = PdfReader(io.BytesIO(file_data))
        num_pages = len(reader.pages)
        all_text = []

        for i, page in enumerate(reader.pages):
            text = page.extract_text()
            if text and text.strip():
                all_text.append(f"--- Sayfa {i+1} ---\n{text.strip()}")

        if not all_text:
            return f"PDF okundu ({num_pages} sayfa) ancak metin çıkarılamadı. Muhtemelen görsel/tabana dayalı PDF."

        content = "\n\n".join(all_text)
        return content[:MAX_TEXT_CONTENT]
    except Exception as e:
        return f"PDF okuma hatası: {e}"


def _count_pdf_pages(file_data: bytes) -> int:
    """Count PDF pages."""
    try:
        from pypdf import PdfReader
        reader = PdfReader(io.BytesIO(file_data))
        return len(reader.pages)
    except Exception:
        return 0


def _extract_docx(file_data: bytes) -> str:
    """Extract text from DOCX using python-docx."""
    try:
        from docx import Document

        doc = Document(io.BytesIO(file_data))
        paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]

        # Also extract tables
        tables_text = []
        for i, table in enumerate(doc.tables[:5]):  # First 5 tables
            rows = []
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells]
                rows.append(" | ".join(cells))
            if rows:
                tables_text.append(f"Tablo {i+1}:\n" + "\n".join(rows))

        parts = []
        if paragraphs:
            parts.append("\n\n".join(paragraphs))
        if tables_text:
            parts.append("\n\n".join(tables_text))

        content = "\n\n---\n\n".join(parts) if parts else "Belge boş."
        return content[:MAX_TEXT_CONTENT]
    except Exception as e:
        return f"DOCX okuma hatası: {e}"


def _extract_xlsx(file_data: bytes, filename: str = "") -> str:
    """Extract data from XLSX using openpyxl."""
    try:
        from openpyxl import load_workbook

        wb = load_workbook(io.BytesIO(file_data), read_only=True, data_only=True)
        all_sheets = []

        for sheet_name in wb.sheetnames[:5]:  # First 5 sheets
            ws = wb[sheet_name]
            rows = []
            for i, row in enumerate(ws.iter_rows(values_only=True)):
                if i >= MAX_TABLE_ROWS:
                    break
                # Convert all values to string
                row_str = [str(c) if c is not None else "" for c in row]
                rows.append(row_str)

            if rows:
                header = rows[0] if rows else []
                data = rows[1:] if len(rows) > 1 else []
                sheet_text = f"Sayfa: {sheet_name}\nSütunlar: {', '.join(header)}\nSatır sayısı: {len(data)}"
                for j, r in enumerate(data[:20]):  # Show first 20 rows
                    sheet_text += f"\n  {j+1}. {' | '.join(r)}"
                if len(data) > 20:
                    sheet_text += f"\n  ... ve {len(data)-20} satır daha."
                all_sheets.append(sheet_text)

        wb.close()
        content = "\n\n".join(all_sheets) if all_sheets else "XLSX dosyası boş."
        return content[:MAX_TEXT_CONTENT]
    except Exception as e:
        return f"XLSX okuma hatası: {e}"


def build_chat_context(attachments: list[dict], user_message: str = "") -> str:
    """Build context string from attachments for the model."""
    if not attachments:
        return user_message

    context_parts = []

    for att in attachments:
        content = att.get("content", "")
        filename = att.get("filename", "unknown")
        file_type = att.get("type", "unknown")
        icon = att.get("icon", "📁")

        if content:
            context_parts.append(f"{icon} DOSYA: {filename}\n{content}")
        elif file_type == "image":
            context_parts.append(f"{icon} GÖRSEL: {filename} (görsel analiz gerekli)")

    if context_parts:
        context_block = "\n\n---\n\n".join(context_parts)
        if user_message:
            return f"{context_block}\n\n---\n\nKullanıcı sorusu: {user_message}"
        return context_block

    return user_message


def build_vision_message(attachment: dict, question: str) -> list[dict]:
    """Build Ollama vision API message for image analysis."""
    b64 = attachment.get("base64", "")
    if not b64:
        return []

    return [
        {
            "role": "user",
            "content": question,
            "images": [b64],
        }
    ]


def process_upload(file_data: bytes, filename: str, source: str = "chat") -> dict[str, Any]:
    """Main entry point: process an uploaded file."""
    # Size check
    if len(file_data) > MAX_FILE_SIZE:
        return {
            "ok": False,
            "error": f"Dosya çok büyük ({len(file_data) / (1024*1024):.1f} MB). Maksimum: {MAX_FILE_SIZE / (1024*1024):.0f} MB",
            "filename": filename,
        }

    # Compute hash for dedup
    file_hash = hashlib.md5(file_data).hexdigest()[:12]

    # Extract content
    result = extract_text_content(file_data, filename)
    result["hash"] = file_hash
    result["source"] = source
    result["uploaded_at"] = datetime.now().isoformat()

    # Save file physically
    try:
        from core.file_manager import save_upload
        import io as _io
        stream = _io.BytesIO(file_data)
        save_result = save_upload(stream, filename, source=source)
        result["saved_path"] = save_result.get("path", "")
    except Exception as e:
        result["saved_error"] = str(e)

    return {"ok": True, "attachment": result}
