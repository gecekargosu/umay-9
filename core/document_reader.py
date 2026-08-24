"""
UMAY Document Reader — Çoklu format belge okuma motoru.

Desteklenen formatlar:
- PDF (.pdf)
- Word (.docx)
- Excel (.xlsx, .xls)
- CSV (.csv)
- TXT (.txt)
- Markdown (.md)

Her okuma işlemi için standart bir dict dönüş formatı kullanılır.
"""
from __future__ import annotations

import csv
import io
import json
import os
import re
import time
from pathlib import Path
from typing import Any

from core.utils.action_logger import eylem_baslat, eylem_hata, eylem_tamamla

# ─── Desteklenen Formatlar ───────────────────────────────────────────────────

SUPPORTED_EXTENSIONS = {
    ".pdf": "pdf",
    ".docx": "word",
    ".xlsx": "excel",
    ".xls": "excel",
    ".csv": "csv",
    ".txt": "text",
    ".md": "markdown",
    ".json": "json",
    ".xml": "xml",
    ".html": "html",
    ".htm": "html",
    ".py": "python",
    ".js": "javascript",
    ".ts": "typescript",
    ".java": "java",
    ".c": "c",
    ".cpp": "cpp",
    ".h": "header",
    ".css": "css",
    ".yaml": "yaml",
    ".yml": "yaml",
    ".toml": "toml",
    ".ini": "ini",
    ".cfg": "config",
    ".log": "log",
    ".sql": "sql",
    ".sh": "shell",
    ".bat": "batch",
    ".ps1": "powershell",
    ".env": "env",
    ".gitignore": "gitignore",
    ".dockerignore": "dockerignore",
}

MAX_CONTENT_CHARS = 50_000  # Maksimum karakter sayısı
MAX_EXCEL_ROWS = 500  # Maksimum Excel satır sayısı
MAX_CSV_ROWS = 500  # Maksimum CSV satır sayısı


# ─── Yardımcı Fonksiyonlar ──────────────────────────────────────────────────

def _detect_encoding(file_path: Path) -> str:
    """Dosya encoding'ini tespit et."""
    try:
        import chardet
        with open(file_path, "rb") as f:
            raw = f.read(10_000)
        result = chardet.detect(raw)
        return result.get("encoding", "utf-8") or "utf-8"
    except ImportError:
        return "utf-8"


def _get_file_type(file_path: Path) -> str:
    """Dosya uzantısına göre tip döndür."""
    ext = file_path.suffix.lower()
    return SUPPORTED_EXTENSIONS.get(ext, "unknown")


def _safe_read_text(file_path: Path, encoding: str = "utf-8") -> str:
    """Güvenli metin okuma (encoding hatası durumunda fallback)."""
    try:
        return file_path.read_text(encoding=encoding)
    except UnicodeDecodeError:
        # Fallback: latin-1 her zaman çalışır
        return file_path.read_text(encoding="latin-1", errors="replace")
    except Exception as e:
        return f"[OKUMA HATASI: {e}]"


def _truncate(content: str, max_chars: int = MAX_CONTENT_CHARS) -> str:
    """İçeriği belirli bir karakter sınırıyla kısalt."""
    if len(content) <= max_chars:
        return content
    return content[:max_chars] + f"\n\n[... Kısaltıldı: {len(content)} toplam karakter, ilk {max_chars} gösteriliyor ...]"


# ─── PDF Okuyucu ─────────────────────────────────────────────────────────────

def read_pdf(file_path: Path, max_pages: int = 50) -> dict[str, Any]:
    """PDF dosyasını oku."""
    try:
        from pypdf import PdfReader
    except ImportError:
        return {"error": "pypdf kütüphanesi yüklü değil. `pip install pypdf` ile yükleyin."}

    try:
        reader = PdfReader(str(file_path))
        total_pages = len(reader.pages)
        pages_to_read = min(total_pages, max_pages)

        text_parts = []
        metadata = {}

        # Metadata
        if reader.metadata:
            metadata = {
                "title": reader.metadata.get("/Title", ""),
                "author": reader.metadata.get("/Author", ""),
                "subject": reader.metadata.get("/Subject", ""),
                "creator": reader.metadata.get("/Creator", ""),
            }

        # Sayfaları oku
        for i in range(pages_to_read):
            page = reader.pages[i]
            page_text = page.extract_text() or ""
            if page_text.strip():
                text_parts.append(f"--- Sayfa {i + 1} ---\n{page_text}")

        content = "\n\n".join(text_parts)

        return {
            "path": str(file_path.name),
            "type": "pdf",
            "total_pages": total_pages,
            "pages_read": pages_to_read,
            "content": _truncate(content),
            "metadata": metadata,
            "char_count": len(content),
            "status": "OK",
        }
    except Exception as e:
        return {"error": f"PDF okuma hatası: {e}", "path": str(file_path.name), "status": "ERROR"}


# ─── Word Okuyucu ────────────────────────────────────────────────────────────

def read_word(file_path: Path) -> dict[str, Any]:
    """Word (.docx) dosyasını oku."""
    try:
        from docx import Document
    except ImportError:
        return {"error": "python-docx kütüphanesi yüklü değil. `pip install python-docx` ile yükleyin."}

    try:
        doc = Document(str(file_path))

        # Metadata
        metadata = {}
        if doc.core_properties:
            metadata = {
                "title": doc.core_properties.title or "",
                "author": doc.core_properties.author or "",
                "subject": doc.core_properties.subject or "",
                "created": str(doc.core_properties.created or ""),
                "modified": str(doc.core_properties.modified or ""),
            }

        # Paragrafları oku
        paragraphs = []
        for para in doc.paragraphs:
            text = para.text.strip()
            if text:
                # Stil bilgisini ekle
                style = para.style.name if para.style else ""
                if style and style != "Normal":
                    paragraphs.append(f"[{style}] {text}")
                else:
                    paragraphs.append(text)

        # Tabloları oku
        tables = []
        for table_idx, table in enumerate(doc.tables):
            table_data = []
            for row in table.rows:
                row_data = [cell.text.strip() for cell in row.cells]
                table_data.append(row_data)
            if table_data:
                tables.append({"table_index": table_idx, "rows": table_data})

        content = "\n\n".join(paragraphs)

        # Tabloları metin olarak da ekle
        if tables:
            content += "\n\n--- TABLOLAR ---\n"
            for table in tables:
                content += f"\nTablo {table['table_index'] + 1}:\n"
                for row in table["rows"]:
                    content += " | ".join(row) + "\n"

        return {
            "path": str(file_path.name),
            "type": "word",
            "content": _truncate(content),
            "metadata": metadata,
            "paragraph_count": len(paragraphs),
            "table_count": len(tables),
            "char_count": len(content),
            "status": "OK",
        }
    except Exception as e:
        return {"error": f"Word okuma hatası: {e}", "path": str(file_path.name), "status": "ERROR"}


# ─── Excel Okuyucu ───────────────────────────────────────────────────────────

def read_excel(file_path: Path, max_rows: int = MAX_EXCEL_ROWS) -> dict[str, Any]:
    """Excel (.xlsx/.xls) dosyasını oku."""
    try:
        from openpyxl import load_workbook
    except ImportError:
        return {"error": "openpyxl kütüphanesi yüklü değil. `pip install openpyxl` ile yükleyin."}

    try:
        wb = load_workbook(str(file_path), read_only=True, data_only=True)
        sheets_data = []
        total_rows = 0

        for sheet_name in wb.sheetnames:
            ws = wb[sheet_name]
            rows = []
            row_count = 0

            for row in ws.iter_rows(values_only=True):
                if row_count >= max_rows:
                    break
                # None değerleri boş string ile değiştir
                row_data = [str(cell) if cell is not None else "" for cell in row]
                if any(cell.strip() for cell in row_data):  # Boş satırları atla
                    rows.append(row_data)
                    row_count += 1

            total_rows += row_count
            if rows:
                sheets_data.append({
                    "sheet_name": sheet_name,
                    "row_count": row_count,
                    "rows": rows,
                })

        wb.close()

        # Metin formatında göster
        content_parts = []
        for sheet in sheets_data:
            content_parts.append(f"--- Sayfa: {sheet['sheet_name']} ({sheet['row_count']} satır) ---")
            for row in sheet["rows"]:
                content_parts.append(" | ".join(row))

        content = "\n".join(content_parts)

        return {
            "path": str(file_path.name),
            "type": "excel",
            "content": _truncate(content),
            "sheets": [{"name": s["sheet_name"], "rows": s["row_count"]} for s in sheets_data],
            "total_rows": total_rows,
            "char_count": len(content),
            "status": "OK",
        }
    except Exception as e:
        return {"error": f"Excel okuma hatası: {e}", "path": str(file_path.name), "status": "ERROR"}


# ─── CSV Okuyucu ─────────────────────────────────────────────────────────────

def read_csv(file_path: Path, max_rows: int = MAX_CSV_ROWS, encoding: str | None = None) -> dict[str, Any]:
    """CSV dosyasını oku."""
    try:
        if encoding is None:
            encoding = _detect_encoding(file_path)

        rows = []
        delimiter = ","

        # Encoding'i test et
        try:
            with open(file_path, encoding=encoding) as f:
                sample = f.read(5_000)
        except UnicodeDecodeError:
            encoding = "utf-8"
            with open(file_path, encoding=encoding, errors="replace") as f:
                sample = f.read(5_000)

        # Delimiter'ı tespit et
        try:
            sniffer = csv.Sniffer()
            dialect = sniffer.sniff(sample)
            delimiter = dialect.delimiter
        except csv.Error:
            delimiter = ","

        # CSV'yi oku
        with open(file_path, encoding=encoding, errors="replace") as f:
            reader = csv.reader(f, delimiter=delimiter)
            for i, row in enumerate(reader):
                if i >= max_rows:
                    break
                rows.append(row)

        # Metin formatında göster
        content_parts = []
        if rows:
            # Header varsa göster
            content_parts.append("BAŞLIK: " + " | ".join(rows[0]))
            content_parts.append("-" * 60)
            for row in rows[1:]:
                content_parts.append(" | ".join(row))

        content = "\n".join(content_parts)

        return {
            "path": str(file_path.name),
            "type": "csv",
            "content": _truncate(content),
            "delimiter": delimiter,
            "encoding": encoding,
            "row_count": len(rows),
            "header": rows[0] if rows else [],
            "char_count": len(content),
            "status": "OK",
        }
    except Exception as e:
        return {"error": f"CSV okuma hatası: {e}", "path": str(file_path.name), "status": "ERROR"}


# ─── Metin Dosyası Okuyucu ───────────────────────────────────────────────────

def read_text(file_path: Path, encoding: str | None = None) -> dict[str, Any]:
    """TXT/Markdown/Code dosyasını oku."""
    try:
        if encoding is None:
            encoding = _detect_encoding(file_path)

        content = _safe_read_text(file_path, encoding)
        file_type = _get_file_type(file_path)

        # Markdown ise başlık analizi yap
        headings = []
        if file_type == "markdown":
            for line in content.split("\n"):
                if line.startswith("#"):
                    headings.append(line.strip())

        return {
            "path": str(file_path.name),
            "type": file_type,
            "content": _truncate(content),
            "encoding": encoding,
            "line_count": len(content.split("\n")),
            "char_count": len(content),
            "headings": headings[:20] if headings else [],
            "status": "OK",
        }
    except Exception as e:
        return {"error": f"Metin okuma hatası: {e}", "path": str(file_path.name), "status": "ERROR"}


# ─── Ana Okuma Fonksiyonu ───────────────────────────────────────────────────

def read_document(file_path: str | Path, **kwargs) -> dict[str, Any]:
    """
    Belge okuma ana fonksiyonu. Dosya tipine göre otomatik olarak
    uygun okuyucuyu seçer.

    Args:
        file_path: Okunacak dosya yolu (mutlak veya relative)
        **kwargs: Format-specific parametreler

    Returns:
        dict: Okunan belge bilgileri
    """
    path = Path(file_path)

    if not path.exists():
        return {"error": f"Dosya bulunamadı: {path}", "status": "ERROR"}

    if not path.is_file():
        return {"error": f"Bu bir dosya değil: {path}", "status": "ERROR"}

    file_type = _get_file_type(path)

    # Action logging
    aid = eylem_baslat(
        ajan="document_reader",
        niyet=f"Belge oku: {path.name}",
        plan=f"Format: {file_type}, Boyut: {path.stat().st_size} byte",
        model="",
    )

    try:
        if file_type == "pdf":
            result = read_pdf(path, max_pages=kwargs.get("max_pages", 50))
        elif file_type == "word":
            result = read_word(path)
        elif file_type == "excel":
            result = read_excel(path, max_rows=kwargs.get("max_rows", MAX_EXCEL_ROWS))
        elif file_type == "csv":
            result = read_csv(path, max_rows=kwargs.get("max_rows", MAX_CSV_ROWS),
                            encoding=kwargs.get("encoding"))
        elif file_type in ("text", "markdown", "json", "xml", "html", "python",
                          "javascript", "typescript", "java", "c", "cpp", "header",
                          "css", "yaml", "toml", "ini", "config", "log", "sql",
                          "shell", "batch", "powershell", "env", "gitignore", "dockerignore"):
            result = read_text(path, encoding=kwargs.get("encoding"))
        else:
            result = {"error": f"Desteklenmeyen dosya formatı: {path.suffix}",
                     "supported": list(SUPPORTED_EXTENSIONS.keys()),
                     "status": "ERROR"}

        if "error" in result:
            eylem_hata(aid, result["error"])
        else:
            eylem_tamamla(
                aid,
                f"{path.name} okundu ({result.get('char_count', 0)} karakter)",
                True,
                0,
            )

        result["file_size"] = path.stat().st_size
        result["absolute_path"] = str(path.resolve())
        return result

    except Exception as e:
        eylem_hata(aid, str(e))
        return {"error": f"Beklenmeyen hata: {e}", "path": str(path.name), "status": "ERROR"}


# ─── Klasör Tarama ──────────────────────────────────────────────────────────

def scan_directory(
    dir_path: str | Path,
    recursive: bool = True,
    file_types: list[str] | None = None,
    max_files: int = 200,
) -> dict[str, Any]:
    """
    Klasörü tara ve desteklenen belgeleri listele.

    Args:
        dir_path: Taranacak klasör yolu
        recursive: Alt klasörleri de tara
        file_types: Sadece belirli tipleri dahil et (örn: ["pdf", "word", "excel"])
        max_files: Maksimum dosya sayısı

    Returns:
        dict: Tarama sonuçları
    """
    root = Path(dir_path)

    if not root.exists() or not root.is_dir():
        return {"error": f"Klasör bulunamadı: {root}", "status": "ERROR"}

    # file_types filtresi
    type_filter = None
    if file_types:
        type_filter = set()
        for ft in file_types:
            for ext, ftype in SUPPORTED_EXTENSIONS.items():
                if ftype == ft:
                    type_filter.add(ext)

    files = []
    skipped_dirs = {".git", "__pycache__", ".venv", "node_modules", "dist", "build",
                   ".next", ".cache", ".umay_backups", "chroma"}

    iterator = root.rglob("*") if recursive else root.iterdir()

    for item in iterator:
        if len(files) >= max_files:
            break

        # Klasörleri atla
        if item.is_dir():
            continue

        # Skip directories
        if any(part in skipped_dirs for part in item.parts):
            continue

        # Tip filtresi
        if type_filter and item.suffix.lower() not in type_filter:
            continue

        # Desteklenen format kontrolü
        file_type = _get_file_type(item)
        if file_type == "unknown" and type_filter is None:
            continue

        try:
            stat = item.stat()
            files.append({
                "path": str(item.relative_to(root)),
                "name": item.name,
                "type": file_type,
                "size": stat.st_size,
                "modified": stat.st_mtime,
            })
        except OSError:
            continue

    # Boyuta göre sırala (büyükten küçüğe)
    files.sort(key=lambda x: x["size"], reverse=True)

    # İstatistikler
    type_counts = {}
    total_size = 0
    for f in files:
        ft = f["type"]
        type_counts[ft] = type_counts.get(ft, 0) + 1
        total_size += f["size"]

    return {
        "directory": str(root),
        "recursive": recursive,
        "file_count": len(files),
        "total_size": total_size,
        "type_counts": type_counts,
        "files": files,
        "status": "OK",
    }


# ─── Belgelerde Arama ───────────────────────────────────────────────────────

def search_in_documents(
    query: str,
    dir_path: str | Path | None = None,
    file_types: list[str] | None = None,
    max_results: int = 50,
) -> dict[str, Any]:
    """
    Belgelerde metin ara.

    Args:
        query: Arama sorgusu (regex destekli)
        dir_path: Arama yapılacak klasör (None ise aktif workspace)
        file_types: Sadece belirli tipleri ara
        max_results: Maksimum sonuç sayısı

    Returns:
        dict: Arama sonuçları
    """
    from core.agent_tools import ACTIVE_WORKSPACE

    root = Path(dir_path) if dir_path else ACTIVE_WORKSPACE

    if not root.exists():
        return {"error": f"Klasör bulunamadı: {root}", "status": "ERROR"}

    # Regex derle
    try:
        pattern = re.compile(query, re.IGNORECASE)
    except re.error:
        # Regex değilse literal ara
        pattern = re.compile(re.escape(query), re.IGNORECASE)

    # file_types filtresi
    type_filter = None
    if file_types:
        type_filter = set()
        for ft in file_types:
            for ext, ftype in SUPPORTED_EXTENSIONS.items():
                if ftype == ft:
                    type_filter.add(ext)

    results = []
    scanned_files = 0
    skipped_dirs = {".git", "__pycache__", ".venv", "node_modules", "dist", "build",
                   ".next", ".cache", ".umay_backups", "chroma"}

    for file_path in root.rglob("*"):
        if len(results) >= max_results:
            break

        if not file_path.is_file():
            continue

        if any(part in skipped_dirs for part in file_path.parts):
            continue

        # Tip filtresi
        if type_filter and file_path.suffix.lower() not in type_filter:
            continue

        # Desteklenen format kontrolü
        file_type = _get_file_type(file_path)
        if file_type == "unknown":
            continue

        scanned_files += 1

        # Dosyayı oku ve ara
        try:
            if file_type in ("pdf",):
                doc = read_pdf(file_path, max_pages=10)
                content = doc.get("content", "")
            elif file_type == "word":
                doc = read_word(file_path)
                content = doc.get("content", "")
            elif file_type == "excel":
                doc = read_excel(file_path, max_rows=50)
                content = doc.get("content", "")
            else:
                content = _safe_read_text(file_path)

            # Satır satır ara
            for lineno, line in enumerate(content.split("\n"), 1):
                if pattern.search(line):
                    results.append({
                        "file": str(file_path.relative_to(root)),
                        "type": file_type,
                        "line": lineno,
                        "text": line.strip()[:500],
                    })
                    if len(results) >= max_results:
                        break
        except Exception:
            continue

    return {
        "query": query,
        "directory": str(root),
        "scanned_files": scanned_files,
        "result_count": len(results),
        "results": results,
        "status": "OK",
    }


# ─── Belge → Hafıza Aktarımı ────────────────────────────────────────────────

def document_to_memory(
    file_path: str | Path,
    source: str = "document",
) -> dict[str, Any]:
    """
    Belge içeriğini RAG/hafıza sistemine aktar.

    Args:
        file_path: Okunacak belge yolu
        source: Kaynak etiketi

    Returns:
        dict: Aktarım sonucu
    """
    try:
        from rag.memory_manager import add_memory
    except ImportError:
        return {"error": "memory_manager modülü bulunamadı.", "status": "ERROR"}

    # Belgeyi oku
    doc = read_document(file_path)

    if "error" in doc:
        return doc

    content = doc.get("content", "")
    if not content.strip():
        return {"error": "Belge içeriği boş.", "path": str(file_path), "status": "ERROR"}

    # İçeriği parçalara böl (ChromaDB için)
    chunk_size = 1000
    chunk_overlap = 200
    chunks = []

    for i in range(0, len(content), chunk_size - chunk_overlap):
        chunk = content[i:i + chunk_size]
        if chunk.strip():
            chunks.append(chunk)

    # Her parçayı hafızaya ekle
    added_count = 0
    for i, chunk in enumerate(chunks):
        metadata = f"{source}:{doc.get('path', 'unknown')}:chunk_{i}"
        if add_memory(chunk, source=metadata):
            added_count += 1

    return {
        "path": str(file_path),
        "type": doc.get("type", "unknown"),
        "total_chunks": len(chunks),
        "added_to_memory": added_count,
        "char_count": len(content),
        "status": "OK",
    }


# ─── Test Fonksiyonu ────────────────────────────────────────────────────────

if __name__ == "__main__":
    print("=== UMAY Document Reader Test ===\n")

    # Test: Bir dosyayı oku
    test_file = Path(__file__).parent.parent / "README.md"
    if test_file.exists():
        result = read_document(test_file)
        print(f"Dosya: {result.get('path')}")
        print(f"Tip: {result.get('type')}")
        print(f"Karakter: {result.get('char_count')}")
        print(f"Durum: {result.get('status')}")
        print(f"İçerik önizleme: {result.get('content', '')[:200]}...")
    else:
        print(f"Test dosyası bulunamadı: {test_file}")
