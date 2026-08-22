"""
UMAY Document Reader Tests
PDF, Word, Excel, CSV, TXT, Markdown okuma testleri.
"""
import csv
import json
import os
from pathlib import Path

import pytest


# ─── Test Fixtures ──────────────────────────────────────────────────────────

@pytest.fixture
def sample_txt(tmp_path):
    """Örnek TXT dosyası oluştur."""
    content = "Merhaba dünya!\nBu bir test dosyasıdır.\nUMAY AI asistanı test ediliyor."
    file = tmp_path / "test.txt"
    file.write_text(content, encoding="utf-8")
    return file


@pytest.fixture
def sample_markdown(tmp_path):
    """Örnek Markdown dosyası oluştur."""
    content = """# UMAY Projesi

## Giriş
Bu proje bir AI asistanıdır.

## Özellikler
- Hafıza sistemi
- Agent desteği
- Browser entegrasyonu

### Detaylı Bilgi
ChromaDB kullanılır.
"""
    file = tmp_path / "test.md"
    file.write_text(content, encoding="utf-8")
    return file


@pytest.fixture
def sample_csv(tmp_path):
    """Örnek CSV dosyası oluştur."""
    file = tmp_path / "test.csv"
    with open(file, "w", encoding="utf-8", newline="") as f:
        writer = csv.writer(f)
        writer.writerow(["İsim", "Yaş", "Şehir"])
        writer.writerow(["Cengiz", "30", "İstanbul"])
        writer.writerow(["Ahmet", "25", "Ankara"])
        writer.writerow(["Ayşe", "28", "İzmir"])
    return file


@pytest.fixture
def sample_json(tmp_path):
    """Örnek JSON dosyası oluştur."""
    data = {
        "proje": "UMAY",
        "surum": "1.0",
        "ozellikler": ["hafiza", "agent", "browser"],
    }
    file = tmp_path / "test.json"
    file.write_text(json.dumps(data, ensure_ascii=False, indent=2), encoding="utf-8")
    return file


@pytest.fixture
def sample_python(tmp_path):
    """Örnek Python dosyası oluştur."""
    content = '''def hello_world():
    """Merhaba dünya fonksiyonu."""
    print("Hello, World!")
    return True

class UMAYAgent:
    """UMAY Agent sınıfı."""
    
    def __init__(self, name):
        self.name = name
    
    def run(self):
        return f"{self.name} çalışıyor"
'''
    file = tmp_path / "test.py"
    file.write_text(content, encoding="utf-8")
    return file


@pytest.fixture
def sample_word(tmp_path):
    """Örnek Word dosyası oluştur."""
    try:
        from docx import Document
        doc = Document()
        doc.add_heading("UMAY Proje Raporu", 0)
        doc.add_paragraph("Bu bir test raporudur.")
        doc.add_paragraph("UMAY AI asistanı geliştirme aşamasındadır.")
        
        # Tablo ekle
        table = doc.add_table(rows=3, cols=2)
        table.cell(0, 0).text = "Özellik"
        table.cell(0, 1).text = "Durum"
        table.cell(1, 0).text = "Hafıza"
        table.cell(1, 1).text = "Tamam"
        table.cell(2, 0).text = "Agent"
        table.cell(2, 1).text = "Geliştirme Aşamasında"
        
        file = tmp_path / "test.docx"
        doc.save(str(file))
        return file
    except ImportError:
        pytest.skip("python-docx yüklü değil")


@pytest.fixture
def sample_excel(tmp_path):
    """Örnek Excel dosyası oluştur."""
    try:
        from openpyxl import Workbook
        wb = Workbook()
        ws = wb.active
        ws.title = "Veriler"
        
        # Header
        ws.append(["İsim", "Yaş", "Şehir"])
        
        # Veriler
        ws.append(["Cengiz", 30, "İstanbul"])
        ws.append(["Ahmet", 25, "Ankara"])
        ws.append(["Ayşe", 28, "İzmir"])
        
        file = tmp_path / "test.xlsx"
        wb.save(str(file))
        return file
    except ImportError:
        pytest.skip("openpyxl yüklü değil")


@pytest.fixture
def sample_pdf(tmp_path):
    """Örnek PDF dosyası oluştur (basit metin)."""
    try:
        from pypdf import PdfWriter
        from io import BytesIO
        
        writer = PdfWriter()
        
        # Basit bir PDF oluştur
        from pypdf.generic import TextStringObject
        from pypdf._page import PageObject
        
        page = PageObject.create_blank_page(width=595, height=842)
        
        # PDF'i kaydet
        file = tmp_path / "test.pdf"
        with open(file, "wb") as f:
            writer.write(f)
        
        return file
    except Exception:
        pytest.skip("PDF oluşturulamadı")


# ─── Text Reader Tests ──────────────────────────────────────────────────────

class TestTextReader:
    """TXT dosyası okuma testleri."""
    
    def test_read_text_file(self, sample_txt):
        """TXT dosyası okunabilmeli."""
        from core.document_reader import read_text
        result = read_text(sample_txt)
        
        assert result["status"] == "OK"
        assert result["type"] == "text"
        assert "Merhaba dünya" in result["content"]
        assert result["char_count"] > 0
        assert result["line_count"] == 3
    
    def test_read_text_encoding_detection(self, sample_txt):
        """Encoding otomatik algılanabilmeli."""
        from core.document_reader import read_text
        result = read_text(sample_txt)
        
        assert result["status"] == "OK"
        assert "encoding" in result


# ─── Markdown Reader Tests ──────────────────────────────────────────────────

class TestMarkdownReader:
    """Markdown dosyası okuma testleri."""
    
    def test_read_markdown_file(self, sample_markdown):
        """Markdown dosyası okunabilmeli."""
        from core.document_reader import read_text
        result = read_text(sample_markdown)
        
        assert result["status"] == "OK"
        assert result["type"] == "markdown"
        assert "# UMAY Projesi" in result["content"]
        assert len(result["headings"]) > 0
    
    def test_markdown_headings_extracted(self, sample_markdown):
        """Markdown başlıkları çıkarılabilmeli."""
        from core.document_reader import read_text
        result = read_text(sample_markdown)
        
        assert result["status"] == "OK"
        assert any("#" in h for h in result["headings"])


# ─── CSV Reader Tests ───────────────────────────────────────────────────────

class TestCSVReader:
    """CSV dosyası okuma testleri."""
    
    def test_read_csv_file(self, sample_csv):
        """CSV dosyası okunabilmeli."""
        from core.document_reader import read_csv
        result = read_csv(sample_csv)
        
        assert result["status"] == "OK"
        assert result["type"] == "csv"
        assert result["row_count"] == 4  # Header + 3 data rows
        assert "İsim" in result["header"]
    
    def test_csv_delimiter_detection(self, sample_csv):
        """CSV delimiter'ı otomatik algılanabilmeli."""
        from core.document_reader import read_csv
        result = read_csv(sample_csv)
        
        assert result["status"] == "OK"
        assert result["delimiter"] == ","


# ─── JSON Reader Tests ──────────────────────────────────────────────────────

class TestJSONReader:
    """JSON dosyası okuma testleri."""
    
    def test_read_json_file(self, sample_json):
        """JSON dosyası okunabilmeli."""
        from core.document_reader import read_text
        result = read_text(sample_json)
        
        assert result["status"] == "OK"
        assert result["type"] == "json"
        assert "UMAY" in result["content"]


# ─── Python Code Reader Tests ───────────────────────────────────────────────

class TestPythonReader:
    """Python kod dosyası okuma testleri."""
    
    def test_read_python_file(self, sample_python):
        """Python dosyası okunabilmeli."""
        from core.document_reader import read_text
        result = read_text(sample_python)
        
        assert result["status"] == "OK"
        assert result["type"] == "python"
        assert "def hello_world" in result["content"]
        assert "class UMAYAgent" in result["content"]


# ─── Word Reader Tests ──────────────────────────────────────────────────────

class TestWordReader:
    """Word dosyası okuma testleri."""
    
    def test_read_word_file(self, sample_word):
        """Word dosyası okunabilmeli."""
        from core.document_reader import read_word
        result = read_word(sample_word)
        
        assert result["status"] == "OK"
        assert result["type"] == "word"
        assert "UMAY Proje Raporu" in result["content"]
        assert result["paragraph_count"] > 0
    
    def test_word_tables_extracted(self, sample_word):
        """Word tabloları çıkarılabilmeli."""
        from core.document_reader import read_word
        result = read_word(sample_word)
        
        assert result["status"] == "OK"
        assert result["table_count"] > 0


# ─── Excel Reader Tests ─────────────────────────────────────────────────────

class TestExcelReader:
    """Excel dosyası okuma testleri."""
    
    def test_read_excel_file(self, sample_excel):
        """Excel dosyası okunabilmeli."""
        from core.document_reader import read_excel
        result = read_excel(sample_excel)
        
        assert result["status"] == "OK"
        assert result["type"] == "excel"
        assert len(result["sheets"]) > 0
        assert result["total_rows"] > 0
    
    def test_excel_sheets_info(self, sample_excel):
        """Excel sayfa bilgileri döndürülebilmeli."""
        from core.document_reader import read_excel
        result = read_excel(sample_excel)
        
        assert result["status"] == "OK"
        assert any(s["name"] == "Veriler" for s in result["sheets"])


# ─── Main Reader Function Tests ─────────────────────────────────────────────

class TestMainReader:
    """Ana read_document fonksiyonu testleri."""
    
    def test_read_document_txt(self, sample_txt):
        """read_document TXT dosyası okuyabilmeli."""
        from core.document_reader import read_document
        result = read_document(sample_txt)
        
        assert result["status"] == "OK"
        assert result["type"] == "text"
        assert "absolute_path" in result
    
    def test_read_document_csv(self, sample_csv):
        """read_document CSV dosyası okuyabilmeli."""
        from core.document_reader import read_document
        result = read_document(sample_csv)
        
        assert result["status"] == "OK"
        assert result["type"] == "csv"
    
    def test_read_document_nonexistent(self, tmp_path):
        """Olmayan dosya okunursa hata dönmeli."""
        from core.document_reader import read_document
        result = read_document(tmp_path / "nonexistent.txt")
        
        assert result["status"] == "ERROR"
        assert "error" in result
    
    def test_read_document_directory(self, tmp_path):
        """Klasör okunursa hata dönmeli."""
        from core.document_reader import read_document
        result = read_document(tmp_path)
        
        assert result["status"] == "ERROR"


# ─── Directory Scanner Tests ────────────────────────────────────────────────

class TestDirectoryScanner:
    """Klasör tarama testleri."""
    
    def test_scan_directory(self, tmp_path, sample_txt, sample_csv):
        """Klasör taranabilmeli."""
        from core.document_reader import scan_directory
        result = scan_directory(tmp_path)
        
        assert result["status"] == "OK"
        assert result["file_count"] >= 2
        assert "text" in result["type_counts"]
        assert "csv" in result["type_counts"]
    
    def test_scan_directory_with_type_filter(self, tmp_path, sample_txt, sample_csv):
        """Tip filtresi çalışabilmeli."""
        from core.document_reader import scan_directory
        result = scan_directory(tmp_path, file_types=["text"])
        
        assert result["status"] == "OK"
        assert result["file_count"] == 1
        assert "csv" not in result["type_counts"]
    
    def test_scan_directory_nonexistent(self, tmp_path):
        """Olmayan klasör taranırsa hata dönmeli."""
        from core.document_reader import scan_directory
        result = scan_directory(tmp_path / "nonexistent")
        
        assert result["status"] == "ERROR"


# ─── Document Search Tests ──────────────────────────────────────────────────

class TestDocumentSearch:
    """Belge arama testleri."""
    
    def test_search_in_documents(self, tmp_path, sample_txt, sample_markdown):
        """Belgelerde arama yapılabilmeli."""
        from core.document_reader import search_in_documents
        result = search_in_documents("UMAY", dir_path=tmp_path)
        
        assert result["status"] == "OK"
        assert result["result_count"] > 0
        assert any("UMAY" in r["text"] for r in result["results"])
    
    def test_search_with_regex(self, tmp_path, sample_python):
        """Regex araması çalışabilmeli."""
        from core.document_reader import search_in_documents
        result = search_in_documents(r"def \w+", dir_path=tmp_path)
        
        assert result["status"] == "OK"
        assert result["result_count"] > 0


# ─── Document to Memory Tests ───────────────────────────────────────────────

class TestDocumentToMemory:
    """Belge → Hafıza aktarım testleri."""
    
    def test_document_to_memory(self, sample_txt):
        """Belge hafızaya aktarılabilmeli."""
        from core.document_reader import document_to_memory
        result = document_to_memory(sample_txt, source="test")
        
        assert result["status"] == "OK"
        # added_to_memory 0 olabilir (aynı içerik zaten hafızada olabilir)
        assert result["total_chunks"] > 0
        assert result["char_count"] > 0
    
    def test_document_to_memory_with_error(self, tmp_path):
        """Hatalı dosya için hata dönmeli."""
        from core.document_reader import document_to_memory
        result = document_to_memory(tmp_path / "nonexistent.txt")
        
        assert result["status"] == "ERROR"


# ─── Tool Integration Tests ─────────────────────────────────────────────────

class TestToolIntegration:
    """Agent tools entegrasyon testleri."""
    
    def test_document_reader_tools_registered(self):
        """Document reader tool'ları kayıtlı olmalı."""
        from core.agent_tools import TOOLS, DISPATCH
        
        tool_names = {t["function"]["name"] for t in TOOLS}
        assert "read_document" in tool_names
        assert "scan_directory" in tool_names
        assert "search_in_documents" in tool_names
        assert "document_to_memory" in tool_names
        
        assert "read_document" in DISPATCH
        assert "scan_directory" in DISPATCH
        assert "search_in_documents" in DISPATCH
        assert "document_to_memory" in DISPATCH
    
    def test_read_document_via_tools(self, tmp_path):
        """read_document tool'u üzerinden okuma yapılabilmeli."""
        from core.agent_tools import read_document
        
        # Test dosyası oluştur
        test_file = tmp_path / "test.txt"
        test_file.write_text("Test içeriği", encoding="utf-8")
        
        result = read_document(str(test_file))
        assert result["status"] == "OK"
    
    def test_scan_directory_via_tools(self, tmp_path):
        """scan_directory tool'u üzerinden tarama yapılabilmeli."""
        from core.agent_tools import scan_directory
        
        # Test dosyası oluştur
        test_file = tmp_path / "test.txt"
        test_file.write_text("Test", encoding="utf-8")
        
        result = scan_directory(str(tmp_path))
        assert result["status"] == "OK"
        assert result["file_count"] >= 1
